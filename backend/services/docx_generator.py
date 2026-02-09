"""
DOCX Generator Service
=======================
Generate Microsoft Word documents from resume data.
"""

import os
from flask import current_app
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class DOCXGenerator:
    """Service for generating DOCX resumes."""

    def __init__(self):
        self.export_folder = current_app.config.get('EXPORT_FOLDER', 'exports')

    def generate(self, resume_data, template='modern'):
        """Generate DOCX from resume data."""
        try:
            doc = Document()
            
            # Set up document properties
            self._setup_document(doc, resume_data, template)
            
            # Add content sections
            self._add_header(doc, resume_data)
            self._add_contact(doc, resume_data)
            self._add_summary(doc, resume_data)
            self._add_experience(doc, resume_data)
            self._add_education(doc, resume_data)
            self._add_skills(doc, resume_data)
            self._add_projects(doc, resume_data)
            self._add_certifications(doc, resume_data)
            
            # Create export directory if needed
            os.makedirs(self.export_folder, exist_ok=True)
            
            # Generate filename
            filename = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(self.export_folder, filename)
            
            # Save document
            doc.save(filepath)
            
            return {
                'success': True,
                'filename': filename,
                'filepath': filepath
            }
            
        except Exception as e:
            import logging
            logging.error(f"DOCX generation error: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def _setup_document(self, doc, resume_data, template):
        """Set up document defaults."""
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

    def _add_header(self, doc, resume_data):
        """Add header section with name and title."""
        name = resume_data.get('full_name', 'Your Name')
        title = resume_data.get('title', 'Professional Title')
        
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(24)
        name_run.font.name = 'Calibri Light'
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(14)
        title_run.font.name = 'Calibri'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()

    def _add_contact(self, doc, resume_data):
        """Add contact information."""
        contact_parts = []
        
        if resume_data.get('email'):
            contact_parts.append(resume_data['email'])
        if resume_data.get('phone'):
            contact_parts.append(resume_data['phone'])
        if resume_data.get('city'):
            contact_parts.append(resume_data['city'])
        if resume_data.get('country'):
            contact_parts.append(resume_data['country'])
        
        social_parts = []
        if resume_data.get('linkedin'):
            social_parts.append(f"LinkedIn: {resume_data['linkedin']}")
        if resume_data.get('github'):
            social_parts.append(f"GitHub: {resume_data['github']}")
        if resume_data.get('portfolio'):
            social_parts.append(f"Portfolio: {resume_data['portfolio']}")
        
        contact = contact_parts + social_parts
        
        if contact:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.font.size = Pt(10)
            contact_para.add_run(' | '.join(contact))
            
            doc.add_paragraph()

    def _add_summary(self, doc, resume_data):
        """Add professional summary."""
        if not resume_data.get('summary'):
            return
        
        self._add_section_title(doc, 'Professional Summary')
        
        summary_para = doc.add_paragraph()
        summary_para.add_run(resume_data['summary'])
        summary_para.font.size = Pt(11)
        
        doc.add_paragraph()

    def _add_experience(self, doc, resume_data):
        """Add work experience section."""
        experiences = resume_data.get('experiences', [])
        if not experiences:
            return
        
        self._add_section_title(doc, 'Work Experience')
        
        for exp in experiences:
            # Job title and company
            header_para = doc.add_paragraph()
            header_para.keep_together = True
            
            title_run = header_para.add_run(exp.get('position', ''))
            title_run.bold = True
            title_run.font.size = Pt(11)
            
            header_para.add_run(' | ')
            
            company_run = header_para.add_run(exp.get('company', ''))
            company_run.italic = True
            company_run.font.size = Pt(11)
            
            # Date
            date_para = doc.add_paragraph()
            date_para.font.size = Pt(10)
            date_para.font.italic = True
            date_para.add_run(f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}")
            
            # Description
            if exp.get('description'):
                desc_para = doc.add_paragraph()
                desc_para.font.size = Pt(10)
                
                # Format as bullet points
                bullets = exp['description'].split('\n')
                for bullet in bullets:
                    bullet = bullet.strip().lstrip('*-•')
                    if bullet:
                        p = doc.add_paragraph(bullet, style='List Bullet')
                        p.font.size = Pt(10)
            
            doc.add_paragraph()

    def _add_education(self, doc, resume_data):
        """Add education section."""
        education = resume_data.get('education', [])
        if not education:
            return
        
        self._add_section_title(doc, 'Education')
        
        for edu in education:
            # Degree
            header_para = doc.add_paragraph()
            
            degree_run = header_para.add_run(edu.get('degree', ''))
            degree_run.bold = True
            degree_run.font.size = Pt(11)
            
            header_para.add_run(' | ')
            
            inst_run = header_para.add_run(edu.get('institution', ''))
            inst_run.italic = True
            inst_run.font.size = Pt(11)
            
            # Date
            if edu.get('start_date') or edu.get('end_date'):
                date_para = doc.add_paragraph()
                date_para.font.size = Pt(10)
                date_para.add_run(f"{edu.get('start_date', '')} - {edu.get('end_date', '')}")
            
            # Description
            if edu.get('description'):
                desc_para = doc.add_paragraph()
                desc_para.font.size = Pt(10)
                desc_para.add_run(edu['description'])
            
            # GPA
            if edu.get('gpa'):
                gpa_para = doc.add_paragraph()
                gpa_para.font.size = Pt(10)
                gpa_para.add_run(f"GPA: {edu['gpa']}")
            
            doc.add_paragraph()

    def _add_skills(self, doc, resume_data):
        """Add skills section."""
        skills = resume_data.get('skills', [])
        if not skills:
            return
        
        self._add_section_title(doc, 'Skills')
        
        # Group skills by category
        by_category = {}
        for skill in skills:
            category = skill.get('category', 'Other')
            if category not in by_category:
                by_category[category] = []
            by_category[category].append(skill)
        
        if by_category:
            # Add categorized skills
            for category, skill_list in by_category.items():
                if category != 'Other':
                    cat_para = doc.add_paragraph()
                    cat_run = cat_para.add_run(f"{category}: ")
                    cat_run.bold = True
                    
                    skill_names = [s.get('name', '') for s in skill_list]
                    cat_run.add_run(', '.join(skill_names))
                    cat_para.font.size = Pt(10)
        else:
            # Add all skills
            skill_names = [s.get('name', '') for s in skills]
            skills_para = doc.add_paragraph()
            skills_para.font.size = Pt(10)
            skills_para.add_run(', '.join(skill_names))
        
        doc.add_paragraph()

    def _add_projects(self, doc, resume_data):
        """Add projects section."""
        projects = resume_data.get('projects', [])
        if not projects:
            return
        
        self._add_section_title(doc, 'Projects')
        
        for proj in projects:
            # Project name
            header_para = doc.add_paragraph()
            name_run = header_para.add_run(proj.get('name', ''))
            name_run.bold = True
            name_run.font.size = Pt(11)
            
            # Description
            if proj.get('description'):
                desc_para = doc.add_paragraph()
                desc_para.font.size = Pt(10)
                desc_para.add_run(proj['description'])
            
            # Technologies
            if proj.get('technologies'):
                tech_para = doc.add_paragraph()
                tech_para.font.size = Pt(10)
                tech_para.font.italic = True
                tech_para.add_run(f"Technologies: {proj['technologies']}")
            
            doc.add_paragraph()

    def _add_certifications(self, doc, resume_data):
        """Add certifications section."""
        certifications = resume_data.get('certifications', [])
        if not certifications:
            return
        
        self._add_section_title(doc, 'Certifications')
        
        for cert in certifications:
            # Certification name
            header_para = doc.add_paragraph()
            name_run = header_para.add_run(cert.get('name', ''))
            name_run.bold = True
            name_run.font.size = Pt(11)
            
            # Organization
            if cert.get('issuing_organization'):
                org_para = doc.add_paragraph()
                org_para.font.size = Pt(10)
                org_para.add_run(cert['issuing_organization'])
            
            # Date
            if cert.get('issue_date'):
                date_para = doc.add_paragraph()
                date_para.font.size = Pt(10)
                date_para.add_run(f"Issued: {cert['issue_date']}")
            
            doc.add_paragraph()

    def _add_section_title(self, doc, title):
        """Add a section title."""
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(37, 99, 235)  # Blue color
        
        # Add underline
        run.font.underline = True
        
        doc.add_paragraph()
